import sys, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
rPr = style.element.get_or_add_rPr()
rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Microsoft YaHei"/>'))

for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Microsoft YaHei'
    hs.font.color.rgb = RGBColor(0x00, 0x6D, 0x33) if lv <= 2 else RGBColor(0x1B, 0x1C, 0x1C)
    hs.font.bold = True
    hs.font.size = Pt({1:22,2:16,3:13}[lv])
    hs.paragraph_format.space_before = Pt({1:24,2:18,3:12}[lv])
    hs.paragraph_format.space_after = Pt({1:12,2:8,3:6}[lv])
    hrPr = hs.element.get_or_add_rPr()
    hrPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Microsoft YaHei"/>'))

def H(doc, text, level=1):
    return doc.add_heading(text, level=level)

def P(doc, text):
    return doc.add_paragraph(text)

def B(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def Q(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="8" w:color="07C160"/></w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0FAF4" w:val="clear"/>'))
    r = p.add_run(text); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x55,0x55,0x55); r.font.italic = True
    return p

def T(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="006D33" w:val="clear"/>')
        c._element.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val)); r.font.size = Pt(10)
            if ri % 2 == 1:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F6F3F2" w:val="clear"/>')
                c._element.get_or_add_tcPr().append(shd)
    doc.add_paragraph()
    return t

# ═══ COVER ═══
for _ in range(6): doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('公众号排版助手'); r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x00,0x6D,0x33); r.bold = True
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('产品设计说明书'); r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x5D,0x5F,0x5F)
doc.add_paragraph()
dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dp.add_run('一款专为微信公众号创作者打造的「一键排版」工具'); r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
for _ in range(4): doc.add_paragraph()
mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = mp.add_run('版本：v2.0  |  日期：2026-08-03  |  状态：MVP 已交付'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
doc.add_page_break()

# ═══ TOC ═══
H(doc, '目录', 1)
toc = [
    ('一、产品概述', '产品定位 · 核心工作流 · 技术栈'),
    ('二、设计思路', '一键排版心智模型 · 三栏布局 · Markdown 数据层'),
    ('三、关键功能设计', 'AI 排版 · 风格模板 · 主题逆向提取 · 字符格式化 · 一键导出'),
    ('四、完整功能矩阵', '已交付功能分解表'),
    ('五、亮点', '6 个差异化设计优势'),
    ('六、不足与改进方向', '功能缺失 · 体验细节 · 健壮性'),
    ('七、技术架构', '前后端架构 + AI 排版与导出数据流'),
]
for title, desc in toc:
    p = doc.add_paragraph(); r = p.add_run(title); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x00,0x6D,0x33)
    p.add_run(f'  ——  {desc}').font.size = Pt(10)
doc.add_page_break()

# ═══ 1 ═══
H(doc, '一、产品概述', 1)
H(doc, '1.1 产品定位', 2)
P(doc, '公众号排版助手是一款面向微信公众号创作者的轻量级排版工具。产品只做一件事：将用户写好的纯文本，一键转化为排版精良、风格统一的公众号文章，排好后可直接复制到公众号后台发布。')
P(doc, '不同于通用 Markdown 编辑器或拖拽式排版工具（如秀米、135编辑器），本产品的核心差异在于：')
B(doc, '用户不需要具备排版知识', '零门槛：')
B(doc, 'AI 自动完成标题层级划分、关键词高亮、引用金句标注、要点卡片生成', '全自动：')
B(doc, '内置 3 套专业排版风格，一键切换', '风格一致：')
B(doc, '生成微信后台完全兼容的内联样式 HTML，复制粘贴后样式零丢失', '微信原生兼容：')

H(doc, '1.2 核心工作流', 2)
P(doc, '产品的心智模型极简，从打开工具到发布完成仅需四步：')
T(doc, ['步骤', '操作', '耗时'], [['①','打开工具，粘贴或撰写文字','< 30 秒'],['②','点击「AI 结构化排版」按钮','1 次点击'],['③','在手机预览器中确认效果，如需调整可选风格模板','< 1 分钟'],['④','点击「复制正文」，粘贴到公众号后台发布','< 10 秒']])

H(doc, '1.3 技术栈', 2)
T(doc, ['层级','技术选型','说明'],[['前端框架','React 19 + TypeScript','组件化 SPA，类型安全'],['样式方案','Tailwind CSS 4','原子化 CSS，极简设计语言'],['构建工具','Vite 6','秒级 HMR，生产构建 < 90KB gzip'],['后端服务','Express + tsx','API 代理 + 静态文件服务'],['AI 引擎','DeepSeek Chat API','结构化排版智能体'],['数据存储','localStorage','前端全量持久化']])
doc.add_page_break()

# ═══ 2 ═══
H(doc, '二、设计思路', 1)
H(doc, '2.1 一键排版心智模型', 2)
P(doc, '产品的核心设计理念是「输入 → 排版 → 发布」的极简闭环。用户只需提供内容，工具负责呈现——排版决策完全交给 AI。')
Q(doc, '设计原则：每一步的用户阻力都应被降到最低。排版不是目的，发布才是。')

H(doc, '2.2 三栏工位式布局', 2)
P(doc, '界面采用固定三栏布局，模拟「写作桌面 → 发布终端」的心智模型：')
T(doc, ['区域','宽度','职责'],[['左侧导航栏','230px 固定','功能导航：项目管理、模板库、主题提取、正在编辑'],['中央编辑区','弹性伸缩','内容编辑画布 + 格式化工具栏'],['右侧预览器','380px 固定','iPhone 形态实时预览 + 可折叠样式编辑面板']])
P(doc, '编辑在左侧，结果在右侧，两者实时同步（< 100ms 延迟）。创作者无需频繁切换到手机检查效果。')

H(doc, '2.3 Markdown 标记驱动的数据层', 2)
P(doc, '所有内容以纯文本存储，字符级格式化通过轻量 Markdown 标记表达：')
T(doc, ['标记','含义','编辑器显示','预览/导出渲染'],[['## 文字','一级标题','## 文字','带装饰色条的 H2 标题'],['**文字**','加粗','**文字**','<strong> 粗体标签'],['*文字*','斜体','*文字*','<em> 斜体标签'],['> 文字','引用块','> 文字','左侧绿条 + 浅底色的 blockquote'],['::: callout','要点卡片','::: callout 标题','圆角边框 + 主题色背景的卡片']])
P(doc, '编辑器 textarea 中显示原始标记，预览器和微信导出管道自动将标记转换为富文本 HTML。这一设计保证了编辑、预览、导出三层的渲染一致性。')
doc.add_page_break()

# ═══ 3 ═══
H(doc, '三、关键功能设计阐述', 1)
H(doc, '3.1 AI 一键排版（核心功能）', 2)
P(doc, 'AI 排版采用「三层漏斗」架构，逐层确保输出质量：')

H(doc, '第一层：提示词策略（文章类型感知）', 3)
P(doc, 'DeepSeek system prompt 先引导 AI 判断文章类型，再按对应策略排版：')
T(doc, ['文章类型','识别特征','排版策略'],[['观点文','论点+论证，观点密集','引言引用 → 2-3章，每段≤2处加粗，1个callout'],['教程文','步骤清晰，操作性强','每步一个H2，步骤内用列表，1个callout总结'],['清单文','罗列式，条目多','列表为主，简短首尾，不强行加callout'],['叙事文','故事线，情感为主','极简排版，1-2个引用点缀'],['科普文','概念+解释','每个概念一个H2，术语加粗，1-2个callout']])

H(doc, '第二层：密度控制（少即是多）', 3)
P(doc, 'AI 天然倾向于过度排版。系统通过严格数值上限约束：')
T(doc, ['元素','密度规则','设计意图'],[['H2 标题','≤ 字数÷300，最多6个','避免碎片化'],['**加粗**','每段≤2处，全文占比≤15%','避免满屏"重点"'],['> 引用','全文≤4处','引用应稀缺才有分量'],['::: callout','全文≤2个','卡片是点睛之笔'],['列表','全文≤3处','不是所有内容都适合列出来']])

H(doc, '第三层：后处理美化器', 3)
P(doc, 'AI 输出后经过 7 条规则修正，消除「AI 味」：')
T(doc, ['#','规则','解决的问题'],[['1','删除空标题/空引用/空列表','AI 偶发脏输出'],['2','合并过短相邻段落（<30字）','碎片化阅读体验'],['3','限制连续引用/callout（≤2个）','视觉疲劳'],['4','拆分过长段落（>200字）','手机端单段过长'],['5','控制标题数量（≤8个）','层级过深'],['6','控制加粗密度（≤15%）','满屏重点=没有重点'],['7','统一列表标记符号','视觉一致性']])

H(doc, '3.2 排版风格模板', 2)
P(doc, '样式面板顶部提供 3 套预设排版风格，一键切换全局参数：')
T(doc, ['模板','主色','标题风格','字体','行高','场景'],[['极简绿色','#07C160','左侧色条','PingFang SC','1.75x','日常推文'],['杂志黑白','#1F2937','底部粗线','Source Serif 4','1.85x','深度长文'],['商务蓝','#2563EB','色块背景','Microsoft YaHei','1.70x','企业号']])
P(doc, '每套模板展示主色色块和描述，当前激活模板有绿色勾选标识。用户可在模板基础上微调单个参数。')

H(doc, '3.3 模板库与主题逆向提取', 2)
P(doc, '模板库内置 6 个模板。主题逆向提取是差异化功能：输入公众号文章 URL → 自动解析排版 DNA：')
B(doc, '颜色直方图统计，取最高频非灰非白颜色', '主配色：')
B(doc, '遍历内联 style 的 font-size / line-height，取众数', '字号行高：')
B(doc, '检测 H2 的 border-left / background-color / border-bottom', '标题装饰类型：')
B(doc, '检测 blockquote 的左边框色和背景色', '引用框样式：')
B(doc, '检测圆角+背景+大 padding 独立容器，推断卡片类型', '强调卡片：')
P(doc, '提取结果生成可复用 Theme 模板，可预览确认后套用或存入模板库。')

H(doc, '3.4 字符级内联格式化', 2)
P(doc, '选中文字 → 点击 B/I 按钮或 Ctrl+B/I → 自动包裹/取消 ** 和 * 标记 → 预览实时反映。AI 排版时自动批量标注关键词。')

H(doc, '3.5 一键导出', 2)
P(doc, '点击「复制正文」→ 生成微信兼容的全内联 style HTML → Clipboard API 写入 text/html + text/plain → 粘贴到公众号后台。严格遵循微信 HTML 白名单，不使用 flexbox/grid/CSS变量/rem/em。')
doc.add_page_break()

# ═══ 4 ═══
H(doc, '四、完整功能矩阵', 1)
T(doc, ['模块','功能','状态'],[['内容编辑','8 种内容块（H1/H2/P/图片/引用/Callout/列表/分割线）','✅'],['内容编辑','字符级内联格式化（加粗/斜体）+ Ctrl+B/I 快捷键','✅'],['内容编辑','AI 一键结构化排版（三层漏斗质控）','✅'],['样式系统','3 套排版风格模板，一键切换','✅'],['样式系统','全局参数调节（配色/字体/字号/行高/段落间距）','✅'],['样式系统','4 种标题装饰风格','✅'],['模板库','6 个内置模板 + 分类筛选 + 收藏','✅'],['模板库','模板套用创建新文章','✅'],['模板库','主题逆向提取（URL 输入/HTML 粘贴）','✅'],['导出','一键复制到微信（内联样式 HTML）','✅'],['项目管理','文章增删改查 + 搜索 + 分类','✅'],['项目管理','localStorage 持久化','✅'],['辅助','草稿保存时间戳（"X 分钟前" + 绿闪反馈）','✅'],['辅助','实时手机预览（iPhone 外壳可切换）','✅'],['辅助','字数统计 + 阅读时长估算','✅']])
doc.add_page_break()

# ═══ 5 ═══
H(doc, '五、亮点', 1)
hl = [
    ('一键排版闭环','粘贴文字 → 点一下按钮 → 排版完成。AI 理解语义后自动完成全部排版决策，用户只需确认和发布。'),
    ('三层 AI 质控漏斗','提示词层做文章类型感知，密度控制层约束 AI 不过度排版，后处理美化器消除 AI 味——三层递进确保输出质量。'),
    ('文章类型感知策略','观点文、教程文、清单文、叙事文、科普文各有不同排版节奏。AI 先判断类型再排版。'),
    ('主题逆向工程','输入公众号文章 URL → 提取排版 DNA → 生成可复用模板。市面上少有。'),
    ('极简三栏 + 实时预览','编辑/预览/样式三栏同屏，任何修改实时反映，无需切手机检查。'),
    ('微信兼容性深度定制','导出 HTML 严格遵循微信内联样式白名单，复制粘贴后样式零丢失。'),
]
for i, (t, d) in enumerate(hl, 1):
    H(doc, f'亮点 {i}：{t}', 3)
    P(doc, d)
doc.add_page_break()

# ═══ 6 ═══
H(doc, '六、不足与改进方向', 1)
H(doc, '6.1 功能缺失', 2)
T(doc, ['问题','影响','建议'],[['无图片上传','只能引用外链 URL','集成图床或本地上传'],['无 Markdown 导入','无法拖入 .md 文件','添加导入按钮'],['无 PDF 导出','只能复制到剪贴板','增加导出 PDF/文本'],['无移动端适配','桌面端专用','平板优先响应式'],['模板不可原位编辑','只能存为模板','添加模板编辑模式'],['无撤销/重做','误操作无法恢复','Command 模式 undo stack']])
H(doc, '6.2 体验细节', 2)
T(doc, ['问题','说明'],[['字数统计不区分中英文','英文场景一个单词计多个字符'],['预览刷新按钮空操作','点击无实际逻辑'],['模板库仅支持分类筛选','不支持关键词搜索'],['无草稿自动保存动画','有时间戳但无保存中状态']])
H(doc, '6.3 性能与健壮性', 2)
T(doc, ['问题','说明'],[['AI API 外部依赖','DeepSeek 不可用时功能失效（已做本地回退）'],['API Key 硬编码','需环境变量化'],['无错误重试','AI 调用失败直接报错'],['零测试覆盖','无单元/集成测试'],['单用户设计','无账号系统/云端同步']])
doc.add_page_break()

# ═══ 7 ═══
H(doc, '七、技术架构', 1)
P(doc, '产品采用前后端分离架构，前端 React SPA，后端 Express 仅做 API 代理，无数据库依赖。')
T(doc, ['层级','组件','职责'],[['前端渲染','React 19 + Tailwind CSS 4','SPA 界面渲染与状态管理'],['前端编辑','ContentCanvas + Toolbar','区块化编辑 + 内联格式化'],['前端预览','MobilePreview + StylePanel','实时移动端渲染 + 样式调节'],['前端导出','wechatExporter.ts','微信兼容 HTML 生成'],['前端 AI','wechatArticleFormatter.ts','AI 输出解析 + 后处理美化'],['前端存储','localStorage','文章/模板/样式全量持久化'],['后端 API','/api/ai/layout-agent','DeepSeek 结构化排版代理'],['后端 API','/api/extract-theme','公众号文章 HTML 抓取']])

H(doc, '7.1 AI 排版数据流', 2)
T(doc, ['阶段','组件','处理'],[['① 触发','FormattingToolbar','点击按钮 → 收集文本 → 拼接'],['② 调用','App.tsx → server.ts','POST → DeepSeek Chat API'],['③ 推理','DeepSeek','类型判断 → 策略排版 → Markdown'],['④ 解析','wechatArticleFormatter','行解析 → ContentBlock[] → 7条后处理'],['⑤ 渲染','React State','setArticles → 编辑器+预览同步刷新']])

H(doc, '7.2 导出数据流', 2)
T(doc, ['阶段','组件','处理'],[['① 触发','FloatingActionToolbar','点击「复制正文」'],['② 生成','wechatExporter.ts','ContentBlock[]+Style → 全内联HTML'],['③ 写入','Clipboard API','text/html + text/plain → 剪贴板'],['④ 发布','公众号后台','粘贴 → 样式完整保留 → 发布']])

out = r"D:\at happen\layout\公众号排版助手-产品设计说明书.docx"
doc.save(out)
print(f"OK: {out} ({os.path.getsize(out)} bytes)")