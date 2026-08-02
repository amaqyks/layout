from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"C:\Users\99349\Desktop\AI 使用说明书_汇报优化版.docx")


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
GRAY = "666666"
LIGHT_GRAY = "F3F6F8"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2EA", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=10.5, bold=False, color=None, align=None, spacing_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.35
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else "333333")
    return p


def add_bullets(doc, items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = add_para(doc, style=style, spacing_after=2)
        p.paragraph_format.left_indent = Cm(0.55 + level * 0.35)
        r = p.add_run(item)
        set_font(r, size=10.5, color="333333")


def add_numbered(doc, items):
    for item in items:
        p = add_para(doc, style="List Number", spacing_after=2)
        p.paragraph_format.left_indent = Cm(0.55)
        r = p.add_run(item)
        set_font(r, size=10.5, color="333333")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, BLUE)
        set_cell_border(cell, color=BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_font(r, size=10.5, bold=True, color=WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            cell.width = Cm(widths[i])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.text = ""
            set_cell_shading(cell, PALE_BLUE if row_index % 2 == 0 else WHITE)
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=10, color="333333")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                cell.width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def apply_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    for name in ("Normal", "Body Text", "List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10.5)
        style.paragraph_format.line_spacing = 1.35
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    add_para(doc, "汇报材料", size=12, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=16)
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=8)
    r = p.add_run("AI 使用说明书")
    set_font(r, size=24, bold=True, color=BLUE)
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=22)
    r = p.add_run("基于 AI 协同开发公众号排版工具的完整实践记录")
    set_font(r, size=14, color="333333")

    add_table(
        doc,
        ["文档定位", "核心内容", "适用场景"],
        [["AI 协同开发实践复盘", "工具链、工作流、阶段成果与问题解决", "项目汇报 / 经验分享 / 产品案例沉淀"]],
        widths=[4.4, 7.2, 4.4],
    )
    add_para(
        doc,
        "本文档围绕“公众号智能排版工具”的开发过程，梳理从产品构思、UI 设计、前后端实现、Agent Skill 构建到代码优化与研发文档沉淀的完整 AI 协作路径。",
        size=11,
        color="333333",
        spacing_after=4,
    )
    doc.add_page_break()


def add_overview(doc):
    add_heading(doc, "一、AI 工具使用概览")
    rows = [
        ("ChatGPT", "产品分析、需求梳理、Agent 架构设计、技术路线规划"),
        ("Stitch", "生成 UI 设计方案，明确页面结构与产品形态"),
        ("Google AI Studio", "将 UI 参考转化为前端页面代码"),
        ("Coze", "将排版能力拆解为可调用的 Agent Skill"),
        ("Antigravity", "完成后端接口、Agent 接入、前后端联调与功能完善"),
        ("Codex", "进行代码审查、项目结构优化和可维护性提升"),
        ("Proto-to-PRD", "将原型与开发过程整理为规范化研发文档"),
    ]
    add_table(doc, ["AI 工具", "主要用途"], rows, widths=[4, 12])

    add_heading(doc, "二、AI 协同开发工作流")
    workflow = [
        ("1", "产品想法", "明确要解决的公众号排版效率与风格复刻问题"),
        ("2", "ChatGPT 产品分析", "完成产品定位、用户需求、功能拆解与技术路线规划"),
        ("3", "Stitch UI 设计", "生成界面方案，沉淀信息架构和用户操作路径"),
        ("4", "Google AI Studio 前端实现", "依据 UI 图片和资源生成前端页面代码"),
        ("5", "Coze 生成 Agent Skill", "将模板解析与 AI 排版能力封装为可调用 Skill"),
        ("6", "Antigravity 全栈开发", "完成接口、数据处理、Agent 调用和联调"),
        ("7", "Codex 代码审查优化", "优化结构、编码、交互能力和数据可靠性"),
        ("8", "Proto-to-PRD 研发文档", "沉淀产品需求、功能模块、技术方案和流程记录"),
    ]
    add_table(doc, ["步骤", "阶段", "产出重点"], workflow, widths=[1.5, 4.3, 10.2])


def add_stage_sections(doc):
    add_heading(doc, "三、具体开发阶段流程")

    add_heading(doc, "1. 使用 ChatGPT 完成产品规划与开发准备", 2)
    add_table(doc, ["项目", "说明"], [
        ("使用工具", "ChatGPT"),
        ("使用目的", "在正式开发前完成产品定位分析、用户需求梳理、功能拆解、Agent 架构设计与技术路线规划。"),
        ("使用方式", "通过多轮对话逐步明确产品方向，并将模糊想法转化为可执行的开发方案。"),
    ], widths=[3.2, 12.8])
    add_para(doc, "产品定位", size=11, bold=True, color=BLUE)
    add_para(doc, "该产品并非简单的 Markdown 转公众号格式工具，而是一个能够理解公众号视觉风格，并通过 AI 复刻公众号排版模板的智能排版工具。")
    add_para(doc, "核心流程：用户输入公众号文章链接 -> AI 解析文章视觉结构 -> 保存模板 -> 用户输入新文章 -> AI 按照模板自动排版。")
    add_para(doc, "Agent 架构", size=11, bold=True, color=BLUE)
    add_table(doc, ["Agent", "职责", "输入", "处理流程", "输出"], [
        ("公众号模板解析 Agent", "从公众号文章中提取可复用排版模板", "公众号文章链接", "获取 HTML 内容，分析标题样式、正文字体、段落间距、图片布局、引导组件、分割线样式和颜色体系", "结构化模板数据"),
        ("AI 排版 Agent", "根据用户文章和已有模板生成符合公众号风格的富文本内容", "用户文章内容、已保存模板", "内容理解、自动分段、标题识别、样式匹配、富文本结构生成", "公众号富文本结构"),
    ], widths=[3.1, 3.5, 2.9, 4.5, 2])

    stages = [
        ("四、使用 Stitch 完成 UI 设计", "Stitch", "生成产品界面，并进一步明确产品形态。", ["产品定位", "用户流程", "功能需求"], "Stitch 不只是完成 UI 生成，同时帮助重新思考产品的信息架构、用户操作流程、页面功能布局和最终呈现形式。"),
        ("五、使用 Google AI Studio 生成前端", "Google AI Studio", "将 Stitch 生成的 UI 设计转化为真实前端代码。", ["UI 图片参考", "HTML 中引用的图片资源", "页面结构需求"], "由 AI Studio 辅助生成前端页面代码，并根据实际效果持续优化。"),
        ("六、使用 Coze 制作 Agent Skill", "Coze", "将公众号排版能力拆解为可调用的 AI Skill。", ["输入", "输出", "核心职责", "工作流程", "约束条件"], "明确告知 Coze Agent，并生成对应 Skill 文件，下载至本地项目使用。"),
        ("七、使用 Antigravity 完成全栈开发", "Antigravity", "完成完整产品功能开发。", ["后端接口开发", "Agent 调用接入", "前后端联调", "数据处理", "功能完善"], "将产品需求、技术方案以及 Agent Skill 输入 Antigravity，由 AI 辅助完成全栈开发。"),
    ]
    for title, tool, purpose, inputs, value in stages:
        add_heading(doc, title, 1)
        add_table(doc, ["项目", "说明"], [
            ("使用工具", tool),
            ("使用目的", purpose),
            ("输入 / 关注内容", "、".join(inputs)),
            ("AI 带来的价值", value),
        ], widths=[3.2, 12.8])


def add_optimization_and_docs(doc):
    add_heading(doc, "八、使用 Codex 进行代码审查与优化")
    add_table(doc, ["项目", "说明"], [
        ("使用工具", "Codex"),
        ("使用目的", "项目后期代码质量优化，提高代码规范性、可维护性和编辑体验。"),
    ], widths=[3.2, 12.8])
    add_para(doc, "完成任务", size=11, bold=True, color=BLUE)
    add_numbered(doc, [
        "删除无用文件，优化项目结构。",
        "统一 UTF-8 编码，解决字符兼容问题。",
        "增加富文本字符级格式化能力。",
        "增加快捷键支持，提高编辑效率。",
        "增加草稿保存时间戳，提高数据可靠性。",
    ])

    add_heading(doc, "九、使用 Proto-to-PRD 生成研发文档")
    add_table(doc, ["项目", "说明"], [
        ("使用工具", "Proto-to-PRD"),
        ("使用目的", "将已有产品原型和开发过程整理为标准研发文档。"),
        ("输出内容", "产品需求说明、功能模块描述、技术实现方案、开发流程记录。"),
        ("文档价值", "帮助项目从个人开发状态转变为规范化产品文档。"),
    ], widths=[3.2, 12.8])


def add_problems_and_summary(doc):
    add_heading(doc, "十、开发过程中遇到的问题与解决方式")
    add_table(doc, ["问题", "解决方式"], [
        ("Claude 账号封禁", "切换回 Antigravity 作为主要开发工具。"),
        ("直接让 AI 生成完整 Web 应用效果较差", "调整开发流程，采用“UI 设计 -> 前端实现 -> Agent 设计 -> 后端开发”的分阶段方式。"),
        ("AI 生成结果容易修改用户原始文本", "增加内容保护约束，明确输入输出规则。"),
        ("单纯代码生成效果有限", "设计 Skill，将复杂能力拆解为 AI 可调用模块。"),
    ], widths=[6.2, 9.8])

    add_heading(doc, "十一、总结：AI 协同开发模式")
    add_para(doc, "本次项目形成了一套从产品构思、设计、开发到优化的 AI 协同开发工作流。")
    add_table(doc, ["阶段", "核心做法"], [
        ("产品思考阶段", "使用 ChatGPT 完成需求分析与架构设计。"),
        ("设计阶段", "使用 Stitch 快速完成产品视觉设计。"),
        ("开发阶段", "使用 AI Studio、Antigravity 完成前后端实现。"),
        ("智能能力建设阶段", "使用 Coze 将业务能力封装为 Agent Skill。"),
        ("优化阶段", "使用 Codex 提升代码质量。"),
        ("文档阶段", "使用 Proto-to-PRD 沉淀研发资料。"),
    ], widths=[5, 11])
    add_para(doc, "最终，该项目沉淀出一套可复用的 AI 协同开发方法：先用 AI 辅助完成产品理解与方案拆解，再通过设计、开发、Agent 能力建设和代码优化逐步落地，最后形成规范化研发文档。")
    add_para(doc, "本文档已完成结构化排版，重点优化了标题层级、流程逻辑、表格展示和技术文档表达方式，更适合作为项目汇报材料或 AI 产品实践案例使用。", bold=True, color=BLUE)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("AI 使用说明书 | 汇报优化版")
    set_font(r, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("基于 AI 协同开发公众号排版工具的完整实践记录")
    set_font(r, size=9, color=GRAY)


def main():
    doc = Document()
    apply_document_styles(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_overview(doc)
    add_stage_sections(doc)
    add_optimization_and_docs(doc)
    add_problems_and_summary(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
